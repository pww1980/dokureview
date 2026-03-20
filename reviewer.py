import os
import json
import uuid
import re
import datetime
import logging
import anthropic

logger = logging.getLogger("dokureview.reviewer")
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH


SEVERITY_COLORS = {
    "red": "FFB3B3",
    "orange": "FFD9A0",
    "green": "C6EFCE",
}

SEVERITY_LABELS = {
    "red": "Kritisch",
    "orange": "Überarbeitungsbedarf",
    "green": "In Ordnung",
}

MODEL = "claude-sonnet-4-5"
MAX_TOKENS = 8192  # actual maximum for claude-sonnet-4-5


def highlight_run(run, hex_color):
    """Sets background color of a run via shading."""
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    rPr.append(shd)


def add_comment(doc, paragraph, run, comment_text, author="Document Reviewer"):
    """Adds a Word comment to a run via direct XML manipulation."""
    # Get or create comments part; element is stored in ._comments_root
    comments_part = _get_or_create_comments_part(doc)
    comments_element = comments_part._comments_root

    # Generate a unique comment ID
    existing_ids = [
        int(c.get(qn("w:id"), 0))
        for c in comments_element.findall(qn("w:comment"))
    ]
    comment_id = str(max(existing_ids, default=-1) + 1)

    # Build w:comment element
    comment_elem = OxmlElement("w:comment")
    comment_elem.set(qn("w:id"), comment_id)
    comment_elem.set(qn("w:author"), author)
    comment_elem.set(
        qn("w:date"), datetime.datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")
    )
    comment_elem.set(qn("w:initials"), "DR")

    comment_para = OxmlElement("w:p")
    comment_run = OxmlElement("w:r")
    comment_rpr = OxmlElement("w:rPr")
    comment_run.append(comment_rpr)
    comment_t = OxmlElement("w:t")
    comment_t.text = comment_text
    comment_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    comment_run.append(comment_t)
    comment_para.append(comment_run)
    comment_elem.append(comment_para)
    comments_element.append(comment_elem)

    # Insert commentRangeStart before the run
    range_start = OxmlElement("w:commentRangeStart")
    range_start.set(qn("w:id"), comment_id)
    run._r.addprevious(range_start)

    # Insert commentRangeEnd after the run
    range_end = OxmlElement("w:commentRangeEnd")
    range_end.set(qn("w:id"), comment_id)
    run._r.addnext(range_end)

    # Insert commentReference after rangeEnd
    ref_run = OxmlElement("w:r")
    comment_ref = OxmlElement("w:commentReference")
    comment_ref.set(qn("w:id"), comment_id)
    ref_run.append(comment_ref)
    range_end.addnext(ref_run)


def _get_or_create_comments_part(doc):
    """Gets or creates the comments part and ensures ._comments_root is set.

    Key insight: CommentsPart is an XmlPart subclass. XmlPart keeps a single
    cached lxml element in ._element and its blob property serialises from it
    automatically on save — no manual sync needed. We must therefore point
    ._comments_root at the SAME object as ._element, not a freshly parsed copy.
    """
    from docx.opc.part import XmlPart
    from docx.opc.packuri import PackURI
    from docx.oxml import parse_xml

    CT_COMMENTS = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
    REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"

    part = doc.part

    # Return existing part (CommentsPart extends XmlPart: ._element is the root)
    try:
        for rel in part.rels.values():
            if rel.reltype == REL_TYPE:
                cp = rel.target_part
                if not hasattr(cp, "_comments_root"):
                    # Point to the real cached element so mutations are saved
                    if hasattr(cp, "_element"):
                        cp._comments_root = cp._element
                    else:
                        cp._comments_root = parse_xml(cp.blob)
                return cp
    except Exception:
        pass

    # Create a new XmlPart: constructor takes (partname, ct, element, package)
    comments_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        '</w:comments>'
    )
    element = parse_xml(comments_xml.encode("utf-8"))
    cp = XmlPart(PackURI("/word/comments.xml"), CT_COMMENTS, element, part.package)
    # XmlPart stores the element as ._element; point our alias at the same object
    cp._comments_root = cp._element
    part.relate_to(cp, REL_TYPE)
    return cp


def _repair_json(s):
    """Best-effort repair of common JSON errors produced by LLMs.

    The most frequent mistake is a missing comma between two consecutive
    key-value pairs inside an object.  Example:

        "severity": "orange"
        "quote": "..."          ← comma missing after previous value

    Strategy: scan line by line; if a line ends with a JSON value
    (closing quote, digit, true/false/null, }, ]) and the next non-empty
    line starts with a quote character (a new key), insert a comma.
    """
    lines = s.split("\n")
    out = []
    # Regex: line that ends with a JSON value (no trailing comma yet)
    ends_value = re.compile(r'(^|.*[^,])\s*(["\d}\]]|true|false|null)\s*$')
    for i, line in enumerate(lines):
        out.append(line)
        stripped = line.rstrip()
        if not ends_value.match(stripped):
            continue
        if stripped.endswith(","):
            continue
        # Find the next non-empty line
        for j in range(i + 1, len(lines)):
            next_stripped = lines[j].strip()
            if not next_stripped:
                continue
            if next_stripped.startswith('"'):
                # Next line is a new key → insert comma
                out[-1] = stripped + ","
            break
    return "\n".join(out)


def call_claude(text, role):
    """Calls Claude API and returns parsed JSON findings."""
    client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))

    prompt = f"""Analysiere das folgende Dokument sorgfältig.

Gib deine Analyse als strukturiertes JSON zurück. Keine weiteren Erklärungen, nur valides JSON.

Format:
{{
  "summary": "Kurze Gesamtbewertung (2-4 Sätze)",
  "findings": [
    {{
      "id": 1,
      "severity": "red" | "orange" | "green",
      "quote": "WÖRTLICHES Zitat aus dem Dokument – exakt so wie es im Text steht, keine Paraphrasen, keine eigenen Formulierungen, kein [...]  (so kurz wie möglich, max. 150 Zeichen, immer zusammenhängender Text aus EINEM Absatz oder einer Tabellenzelle)",
      "comment": "Erläuterung des Befundes und Handlungsempfehlung"
    }}
  ]
}}

Wichtige Regeln für das Feld 'quote':
- Nur wörtlicher Text der tatsächlich so im Dokument vorkommt
- Kein [...]  keine Auslassungen, kein Kürzen mit Ellipsen
- Keine eigenen Formulierungen oder Zusammenfassungen als Quote
- Wenn kein passender Textteil vorhanden ist: verwende den nächstliegenden konkreten Satz
- Für Tabellen: Text einer einzelnen Zelle oder Zeile verwenden

Severity-Bedeutung:
- red: Kritisches Problem, sofortiger Handlungsbedarf
- orange: Problematisch, sollte überarbeitet werden
- green: Positiv hervorzuheben oder nur geringes Risiko

Dokument:
---
{text}
---"""

    logger.info("Claude-Aufruf: Modell=%s, Rolle='%s', Textlänge=%d Zeichen", MODEL, role.get("name"), len(text))

    for attempt in range(2):
        if attempt > 0:
            logger.warning("Claude-Retry (Versuch %d/2)", attempt + 1)

        response = client.messages.create(
            model=MODEL,
            max_tokens=MAX_TOKENS,
            system=role["system_prompt"],
            messages=[{"role": "user", "content": prompt}],
        )

        raw = response.content[0].text.strip()
        logger.debug("Claude-Antwort (%d Zeichen): %s...", len(raw), raw[:100])

        # Extract JSON: try fenced code block first, then bare JSON object
        json_match = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", raw)
        if json_match:
            raw = json_match.group(1)
        else:
            # Handle truncated responses where closing ``` is missing
            bare_match = re.search(r"(\{[\s\S]*\})", raw)
            if bare_match:
                raw = bare_match.group(1)
                logger.warning("Kein geschlossener Code-Block gefunden, verwende rohen JSON-Block")

        try:
            result = json.loads(raw)
            if "summary" in result and "findings" in result:
                logger.info("Claude-Antwort erfolgreich geparst: %d Findings", len(result["findings"]))
                return result
        except json.JSONDecodeError as e:
            logger.warning("JSON-Parse-Fehler (Versuch %d): %s", attempt + 1, e)
            # Attempt repair before giving up or retrying.
            # LLMs occasionally omit commas between key-value pairs.
            repaired = _repair_json(raw)
            if repaired != raw:
                try:
                    result = json.loads(repaired)
                    if "summary" in result and "findings" in result:
                        logger.info(
                            "JSON nach Reparatur erfolgreich geparst: %d Findings",
                            len(result["findings"]),
                        )
                        return result
                except json.JSONDecodeError:
                    pass
            if attempt == 1:
                raise ValueError(f"Claude hat kein valides JSON zurückgegeben: {raw[:200]}")
            continue
        except KeyError as e:
            logger.warning("JSON-Parse-Fehler (Versuch %d): %s", attempt + 1, e)
            if attempt == 1:
                raise ValueError(f"Claude hat kein valides JSON zurückgegeben: {raw[:200]}")
            continue

    raise ValueError("Claude-Analyse fehlgeschlagen")


def find_quote_in_paragraphs(paragraphs, quote):
    """
    Fuzzy search for quote in paragraphs.
    Returns list of (para_index, start_char, end_char) tuples.
    """
    norm_quote = _norm(quote)
    if not norm_quote:
        return []

    results = []
    for i, para in enumerate(paragraphs):
        norm_text = _norm(para.text)
        idx = norm_text.find(norm_quote)
        if idx != -1:
            results.append((i, idx, idx + len(norm_quote)))
            break  # Use first match

    return results


def split_run_at_positions(para, start, end):
    """
    Splits runs in paragraph to isolate characters from start to end.
    Returns the run covering [start, end).
    """
    # Build character map: list of (run_index, char_in_run)
    runs = list(para.runs)
    char_map = []
    for ri, run in enumerate(runs):
        for ci in range(len(run.text)):
            char_map.append((ri, ci))

    if not char_map or start >= len(char_map) or end > len(char_map):
        return None

    # Find which runs we need to split
    start_run_idx, start_char_idx = char_map[start]
    end_run_idx, end_char_idx = char_map[end - 1]

    # Split end run first (to not mess up indices)
    if end_char_idx < len(runs[end_run_idx].text) - 1:
        _split_run(para, end_run_idx, end_char_idx + 1)
        runs = list(para.runs)
        # Rebuild char_map after split
        char_map = []
        for ri, run in enumerate(runs):
            for ci in range(len(run.text)):
                char_map.append((ri, ci))
        start_run_idx, start_char_idx = char_map[start]
        end_run_idx, end_char_idx = char_map[end - 1]

    # Split start run
    if start_char_idx > 0:
        _split_run(para, start_run_idx, start_char_idx)
        runs = list(para.runs)
        char_map = []
        for ri, run in enumerate(runs):
            for ci in range(len(run.text)):
                char_map.append((ri, ci))
        start_run_idx, _ = char_map[start]
        end_run_idx, _ = char_map[end - 1]

    # If the quote spans multiple runs, merge them into one
    runs = list(para.runs)
    if start_run_idx == end_run_idx:
        return runs[start_run_idx]

    # Merge runs from start_run_idx to end_run_idx
    first_run = runs[start_run_idx]
    merged_text = "".join(r.text for r in runs[start_run_idx: end_run_idx + 1])
    first_run.text = merged_text
    for r in runs[start_run_idx + 1: end_run_idx + 1]:
        r._r.getparent().remove(r._r)

    return first_run


def _split_run(para, run_idx, split_pos):
    """Splits a run at split_pos, inserting a new run after it."""
    runs = list(para.runs)
    run = runs[run_idx]
    original_text = run.text
    before = original_text[:split_pos]
    after = original_text[split_pos:]

    run.text = before

    # Create new run with same formatting
    new_r = OxmlElement("w:r")
    # Copy rPr if exists
    rPr = run._r.find(qn("w:rPr"))
    if rPr is not None:
        import copy
        new_r.append(copy.deepcopy(rPr))
    new_t = OxmlElement("w:t")
    new_t.text = after
    new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    new_r.append(new_t)
    run._r.addnext(new_r)


def apply_finding_to_paragraph(doc, out_para, finding, quote):
    """Highlights and comments the quote in the output paragraph."""
    color = SEVERITY_COLORS.get(finding.get("severity", "green"), "C6EFCE")
    comment_text = f"[{SEVERITY_LABELS.get(finding.get('severity','green'), '')}] {finding.get('comment', '')}"

    actual_text = out_para.text

    # Build regex that allows any whitespace between words (handles collapsed spaces)
    words = re.sub(r"\s+", " ", quote.strip()).split()
    if not words:
        return False
    pattern = r"\s+".join(re.escape(w) for w in words)
    m = re.search(pattern, actual_text, re.IGNORECASE)
    if not m:
        return False

    start, end = m.start(), m.end()

    target_run = split_run_at_positions(out_para, start, end)
    if target_run is None:
        return False

    highlight_run(target_run, color)
    try:
        add_comment(doc, out_para, target_run, comment_text)
    except Exception as e:
        logger.warning("Kommentar konnte nicht eingefügt werden: %s", e)

    return True


def _prepend_summary(out_doc, summary):
    """Insert review summary + divider before the first content in the document body."""
    body = out_doc.element.body
    # Find the first child element (w:p, w:tbl, w:sectPr, …)
    first = body[0] if len(body) > 0 else None

    def _make_p(text=None, style_id=None, bg_color=None, text_color=None):
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        if style_id:
            ps = OxmlElement("w:pStyle")
            ps.set(qn("w:val"), style_id)
            pPr.append(ps)
        if bg_color:
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), bg_color)
            pPr.append(shd)
        p.append(pPr)
        if text:
            r = OxmlElement("w:r")
            if text_color:
                rPr = OxmlElement("w:rPr")
                col = OxmlElement("w:color")
                col.set(qn("w:val"), text_color)
                rPr.append(col)
                r.append(rPr)
            t = OxmlElement("w:t")
            t.text = text
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            r.append(t)
            p.append(r)
        return p

    # Build elements in desired reading order
    elements = [
        _make_p("Review-Zusammenfassung", style_id="Heading1", text_color="1A3A5C"),
        _make_p(summary, bg_color="F2F2F2"),
        _make_p(),  # spacer
        _make_p("Dokument", style_id="Heading2", text_color="1A3A5C"),
    ]

    if first is not None:
        # Insert in reverse order, updating the reference each time so that
        # each new element is placed before the previously inserted one.
        # Result: elements appear in their original list order before `first`.
        ref = first
        for elem in reversed(elements):
            ref.addprevious(elem)
            ref = elem
    else:
        for elem in elements:
            body.append(elem)


def add_summary_section(out_doc, summary):
    """Adds the summary block at the beginning of the document."""
    heading = out_doc.add_heading("Review-Zusammenfassung", level=1)
    heading.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

    para = out_doc.add_paragraph(summary)
    # Light gray background
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)

    out_doc.add_paragraph()  # spacer


def add_findings_appendix(out_doc, findings):
    """Appends a findings summary table at the end of the document."""
    out_doc.add_page_break()
    heading = out_doc.add_heading("Anhang: Review-Findings", level=1)
    heading.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

    table = out_doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Schweregrad"
    hdr_cells[1].text = "Textstelle"
    hdr_cells[2].text = "Kommentar"

    for cell in hdr_cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    # Data rows
    for finding in findings:
        row_cells = table.add_row().cells
        severity = finding.get("severity", "green")
        row_cells[0].text = SEVERITY_LABELS.get(severity, severity)
        row_cells[1].text = finding.get("quote", "")[:200]
        row_cells[2].text = finding.get("comment", "")

        # Color the severity cell
        color = SEVERITY_COLORS.get(severity, "C6EFCE")
        tc = row_cells[0]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color)
        tcPr.append(shd)


def _norm(s):
    """Normalise text for fuzzy matching.

    Handles:
    - Non-breaking and other exotic spaces (\u00a0, \u2009, etc.)
    - Table-extraction separators (' | ') added during text extraction
    - Claude's ellipsis notation ([...] / […]) produced when it truncates
    - Standard whitespace collapsing + lower-case
    """
    # Exotic whitespace → regular space
    s = re.sub(r"[\u00a0\u2009\u202f\u2002\u2003\u2060]", " ", s)
    # Table-extraction separators (added in app.py when joining cells)
    s = re.sub(r"\s*\|\s*", " ", s)
    # Claude ellipsis markers
    s = re.sub(r"\[\.\.\.?\]|\u2026", " ", s)
    # Collapse whitespace, strip, lowercase
    return re.sub(r"\s+", " ", s.strip()).lower()


def _collect_all_paragraphs(doc):
    """Return (text, paragraph) tuples for ALL paragraphs in the document,
    including those inside table cells.  Plain doc.paragraphs skips tables."""
    from docx.oxml.ns import qn as _qn
    from docx.text.paragraph import Paragraph

    result = []
    for block in doc.element.body:
        tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag
        if tag == "p":
            p_obj = Paragraph(block, doc)
            txt = p_obj.text
            if txt.strip():
                result.append((txt, p_obj))
        elif tag == "tbl":
            for p_elem in block.iter(_qn("w:p")):
                p_obj = Paragraph(p_elem, doc)
                txt = p_obj.text
                if txt.strip():
                    result.append((txt, p_obj))
    return result


def _find_table_paragraph_for_quote(doc, quote):
    """Search for a quote across the joined text of each table row.

    Claude often cites text that spans multiple cells or comes from a table
    formatted as 'cell1 | cell2'.  This function joins all cell texts of each
    row and returns the first paragraph of the first row whose combined text
    contains (a prefix of) the quote.
    """
    from docx.oxml.ns import qn as _qn
    from docx.text.paragraph import Paragraph

    norm_quote = _norm(quote)
    # Use progressively shorter prefixes so we get *some* match
    min_len = max(15, len(norm_quote) // 3)
    prefixes = []
    for frac in (1.0, 0.6, 0.4):
        chunk = norm_quote[: max(min_len, int(len(norm_quote) * frac))]
        if chunk not in prefixes:
            prefixes.append(chunk)

    for table in doc.tables:
        for row in table.rows:
            seen = set()
            first_para = None
            texts = []
            for cell in row.cells:
                if id(cell) in seen:
                    continue
                seen.add(id(cell))
                for para in cell.paragraphs:
                    t = para.text.strip()
                    if t:
                        texts.append(t)
                        if first_para is None:
                            first_para = para
            if first_para is None:
                continue
            row_norm = _norm(" ".join(texts))
            for prefix in prefixes:
                if prefix in row_norm:
                    return first_para
    return None


def process_document(text, role, output_dir, source_doc=None):
    """
    Main entry point: takes document text and role, returns path to output DOCX.
    If source_doc (python-docx Document) is provided, paragraph formatting is preserved.
    """
    logger.info("Starte Dokumentenverarbeitung: %d Zeichen, Rolle='%s'", len(text), role.get("name"))

    # Call Claude
    analysis = call_claude(text, role)
    summary = analysis.get("summary", "")
    findings = analysis.get("findings", [])
    logger.info("Analyse erhalten: %d Findings", len(findings))

    unmatched_findings = []
    output_paragraphs = []

    if source_doc is not None:
        # Clone the entire source document via BytesIO so all relationships,
        # styles and numbering definitions remain valid in the output file.
        from io import BytesIO
        buf = BytesIO()
        source_doc.save(buf)
        buf.seek(0)
        out_doc = Document(buf)

        # Collect paragraph references BEFORE inserting the summary header,
        # so we only match against real document content.
        # Include table-cell paragraphs so quotes from tables can be matched.
        output_paragraphs = _collect_all_paragraphs(out_doc)

        # Insert summary + divider before the first content element.
        _prepend_summary(out_doc, summary)
    else:
        # Plain text / TXT input: build a fresh document.
        out_doc = Document()
        add_summary_section(out_doc, summary)
        divider = out_doc.add_heading("Dokument", level=2)
        if divider.runs:
            divider.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

        for para_text in [p.strip() for p in text.split("\n") if p.strip()]:
            out_para = out_doc.add_paragraph(para_text)
            output_paragraphs.append((para_text, out_para))

    # Apply findings
    for finding in findings:
        quote = finding.get("quote", "").strip()
        if not quote:
            continue

        matched = False
        for _para_text, out_para in output_paragraphs:
            if apply_finding_to_paragraph(out_doc, out_para, finding, quote):
                matched = True
                break

        if not matched and source_doc is not None:
            # Fallback: search across joined table-row text.  Claude often
            # cites text that spans multiple cells; individual cell paragraphs
            # won't match the full quote.
            tbl_para = _find_table_paragraph_for_quote(out_doc, quote)
            if tbl_para is not None:
                severity = finding.get("severity", "green")
                color = SEVERITY_COLORS.get(severity, "C6EFCE")
                comment_text = (
                    f"[{SEVERITY_LABELS.get(severity, '')}] {finding.get('comment', '')}"
                )
                if tbl_para.runs:
                    target_run = tbl_para.runs[0]
                else:
                    target_run = tbl_para.add_run("")
                highlight_run(target_run, color)
                try:
                    add_comment(out_doc, tbl_para, target_run, comment_text)
                except Exception as e:
                    logger.warning("Kommentar in Tabellenzelle fehlgeschlagen: %s", e)
                matched = True
                logger.info("Finding #%s per Tabellenzeilen-Fallback platziert",
                            finding.get("id", "?"))

        if not matched:
            logger.warning("Quote nicht im Dokument gefunden (Finding #%s): '%s...'",
                           finding.get("id", "?"), quote[:60])
            unmatched_findings.append(finding)

    # Append unmatched findings as footnotes at the end
    if unmatched_findings:
        out_doc.add_page_break()
        fn_heading = out_doc.add_heading("Nicht zugeordnete Befunde", level=2)
        fn_heading.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

        for finding in unmatched_findings:
            severity = finding.get("severity", "green")
            color = SEVERITY_COLORS.get(severity, "C6EFCE")
            label = SEVERITY_LABELS.get(severity, severity)

            fn_para = out_doc.add_paragraph()
            fn_para.add_run(f"[{label}] ").bold = True
            fn_para.add_run(f'Textstelle: „{finding.get("quote", "")[:200]}" — ')
            fn_para.add_run(finding.get("comment", ""))

            pPr = fn_para._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), color)
            pPr.append(shd)

    # Append findings appendix table
    add_findings_appendix(out_doc, findings)

    if unmatched_findings:
        logger.warning("%d Findings konnten nicht zugeordnet werden", len(unmatched_findings))

    # Save output
    output_path = os.path.join(output_dir, f"review_{uuid.uuid4()}.docx")
    out_doc.save(output_path)
    logger.info("Output-DOCX gespeichert: %s", output_path)
    return output_path
